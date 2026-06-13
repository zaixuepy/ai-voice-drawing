from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局默认样式 ──────────────────────────────────
CHINESE_FONT = 'PingFang SC'      # 正文用
CODE_FONT    = 'Courier New'      # 代码块用
BODY_SIZE    = Pt(11)
CODE_SIZE    = Pt(9)
TITLE_SIZE   = Pt(22)
H1_SIZE      = Pt(16)
H2_SIZE      = Pt(14)
H3_SIZE      = Pt(12)

def set_font(run, name=CHINESE_FONT, size=BODY_SIZE, bold=False, color=None):
    """在每个 run 上强制设置字体（中英文 fallback）"""
    run.font.name = name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # 确保中文字体也生效
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)

def add_h0(text):
    """文档大标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=TITLE_SIZE, bold=True)
    return p

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=H1_SIZE, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))
    p.paragraph_format.outline_level = 0
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=H2_SIZE, bold=True, color=RGBColor(0x33, 0x33, 0x33))
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=H3_SIZE, bold=True, color=RGBColor(0x55, 0x55, 0x55))
    return p

def add_p(text):
    """普通正文段落，支持 \n 换行"""
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        set_font(run)
    return p

def add_bold(text):
    """粗体段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, bold=True)
    return p

def add_italic(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=Pt(10))
    run.italic = True
    return p

def add_code(text):
    """等宽代码块，逐行输出，不额外引入语法高亮"""
    for line in text.split('\n'):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = Cm(0.8)
        run = p.add_run(line if line else ' ')
        set_font(run, name=CODE_FONT, size=CODE_SIZE)
    # 代码块后留一点空
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)

def add_sep():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('─' * 50)
    set_font(run, size=Pt(8), color=RGBColor(0xCC, 0xCC, 0xCC))

def add_table(headers, rows):
    """轻量表格，所有文字统一字体"""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    # 表头
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        set_font(run, size=Pt(10), bold=True)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val)
            set_font(run, size=Pt(10))
    doc.add_paragraph()  # spacing

# ══════════════════════════════════════════════════════
# 正文开始
# ══════════════════════════════════════════════════════

add_h0('VoiceDraw 系统提示词架构说明')

add_italic('本文档整合了两部分内容：第一部分为当前提示词体系的架构说明，第二部分为 Claude Code 给出的提示词工程改进建议。两部分内容均原样保留，未做任何修改。')
doc.add_paragraph()

# ═══════════════════ PART 1 ═══════════════════
doc.add_page_break()
add_h1('第一部分：当前提示词体系全貌')

add_h2('一、管线总览：5 个文件，一条链路')
add_p('用户的语音指令经过 5 层处理变成画布上的图形：')
add_bold('语音 → AppPage.jsx → useCommandProcessor.js → promptBuilder.js → deepseekApi.js → chat.js → DeepSeek API')
add_p('                                                              ↑')
add_p('                                                     canvasSerializer.js (上下文)')

add_table(
    ['文件', '职责', '角色类比'],
    [
        ['promptBuilder.js', '系统提示词 + 用户消息模板', '策略层 — 定义 AI 的行为规则'],
        ['canvasSerializer.js', '将画布元素序列化为中文描述', '感知层 — 让 AI 看到当前画布'],
        ['deepseekApi.js', 'SSE 流式调用 + JSON 提取解析', '容错层 — 确保拿到合法 JSON'],
        ['useCommandProcessor.js', '编排管线 + 失败重试', '编排层 — 协调调用、重试、状态'],
        ['chat.js (backend)', 'API 代理', '传输层 — 转发请求到 DeepSeek'],
    ]
)

# ── 逐文件详解 ──
add_h2('二、逐文件详解')

add_h3('1. promptBuilder.js — 系统提示词（策略层）')
add_p('这是整个项目的核心。当前是 6 层架构，借鉴了 Claude Code 的提示词设计原则：')

add_table(
    ['层级', '内容', '解决的问题'],
    [
        ['第1层：身份', '"你是 VoiceDraw，运行在 Excalidraw 画布上的 AI 绘图助手"', '锚定角色边界，防止模型扮演无关角色'],
        ['第2层：核心原则', '6 条硬约束，用"最高优先级""违反将导致系统崩溃"等语言强调', '输出必须 <output> 包裹、JSON 格式、能执行就不确认、复杂图表≥4 节点'],
        ['第3层：决策框架', '6 级判断树：闲聊→不支持→不完整→矛盾→画布操作→明确指令', '意图识别：何时 execute vs clarify，避免过度确认'],
        ['第4层：输出格式', 'execute/clarify 两种 JSON schema + actions 类型定义', '约束输出结构，让代码能精确提取'],
        ['第5层：布局规则', '坐标范围、排列间距、思维导图固定坐标、流程图模板', '控制元素位置（当前 LLM 空间推理有限）'],
        ['第6层：自检清单', '6 条输出前检查项', '降低 JSON 格式错误率'],
    ]
)

add_p('另外 buildUserMessage 构建用户消息，格式极简：')
add_code('## 当前画布状态\n[画布元素序列化文本]\n## 用户指令\n{用户说的话}')
add_p('设计要点：prompt 中的示例直接包含 <output>{"type":"execute",...}</output> 格式，让模型学习"输出就在标签内"的模式。')

add_h3('2. canvasSerializer.js — 画布上下文（感知层）')
add_p('将 Excalidraw 元素转为一行中文描述，塞进用户消息的 当前画布状态 字段。')
add_code('画布元素：[a1:矩形@(200,200),红色,120x60,label=开始] [a2:箭头,from=a1,to=a3]')
add_p('支持的类型：矩形、椭圆、菱形、箭头、文字、线条。颜色有中英文映射表。')
add_p('作用：让 AI 知道画布上当前有什么——修改、删除、连线都需要感知现有元素。')

add_h3('3. deepseekApi.js — API 调用 + JSON 提取（容错层）')
add_p('两个核心函数：')
add_bold('callDeepSeekApi — SSE 流式调用')
add_p('- 只提取 parsed.choices[0].delta.content（忽略 reasoning_content）\n- 跳过 [DONE] 标记\n- 满 300 字时打印 debug 日志')
add_bold('parseAIResponse — JSON 提取 + 校验，4 层回退：')
add_code('级别1: <output>...</output> 标签提取（优先，prompt 层面引导的）\n级别2: 基础清洗（去 ```json 代码块）\n级别3: tryExtractJSON — 4 种子策略\n    ├── 3a: 字符串感知正向括号匹配（从第一个 { 找配对的 }）\n    ├── 3b: 反向括号匹配（从最后一个 } 反向找 {）\n    ├── 3c: 正则匹配所有 { ... } 候选\n    └── 3d: 从 "type":"execute|clarify" 关键字附近搜索\n级别4: 字段校验（type、actions、question、options）')
add_p('关键亮点 — findMatchingBrace 是字符串感知的，能正确处理：')
add_p('- JSON 字符串内的 { } 不计入深度\n- 转义的 \\" 不误判为字符串边界')
add_p('这是解决 "AI 返回的 JSON 括号不匹配" 问题的核心修复。')

add_h3('4. useCommandProcessor.js — 命令处理编排（编排层）')
add_p('职责：')
add_p('1. 本地命令检测（撤销、清空）— 不走 API，0 延迟\n2. 构造 messages 数组（system prompt + user message）\n3. 调用 callDeepSeekApi + parseAIResponse\n4. 失败自动重试（最多 2 次），重试时追加 JSON 格式提醒\n5. 根据解析结果执行不同分支：execute → 调用 executeActions 写画布，clarify → 返回选项\n6. 维护 chatHistory（最近 10 条）')
add_code('const MAX_RETRIES = 2;\nfor (let attempt = 0; attempt <= MAX_RETRIES; attempt++) {\n    rawResponse = await callDeepSeekApi(...);\n    parsed = parseAIResponse(rawResponse);  // 失败抛异常\n    // ...\n}')

add_h3('5. chat.js (backend) — API 代理（传输层）')
add_p('纯转发 SSE 流，参数：')
add_p('- model: deepseek-v4-flash\n- stream: true\n- temperature: 0.3\n- max_tokens: 2048')
add_p('关键决策：去掉了 response_format: json_object，因为它与 DeepSeek v4 flash 的流式模式不兼容（会导致所有 chunk 的 content: null）。')

# ── 当前架构优点 ──
add_h2('三、当前架构的优点')
add_p('1. 分层清晰：每个文件职责单一，策略/容错/编排/传输互不耦合\n2. 多层容错：JSON 提取 4 层回退 + 最多 3 次 API 重试，单点失败不崩溃\n3. 系统提示词结构化：6 层设计让模型更容易遵循，不是 190 行的"一段话"\n4. 字符串感知括号匹配：正确处理 JSON 字符串内含 {} 和转义引号')

# ── 改进建议 ──
add_h2('四、你还能做的改进（代码层面控制）')

add_h3('改进一：前端自动布局 — 去掉 LLM 的坐标生成职责')
add_p('当前问题：LLM 的 prompt 里塞了大量固定坐标（4 分支坐标、6 分支坐标、间距规则），但 LLM 空间推理能力有天花板，思维导图仍会错位。')
add_p('方案：让 LLM 只输出 节点+关系，坐标由前端算法计算。')
add_code('// 新的 prompt 格式：不再要求坐标\n{"type":"execute","actions":[\n  {"type":"add_node","label":"经济学","shape":"ellipse","level":0},\n  {"type":"add_node","label":"微观经济学","shape":"rect","level":1,"parent":"center"},\n  {"type":"add_node","label":"宏观经济学","shape":"rect","level":1,"parent":"center"},\n  // ...\n]}\n\n// 前端自动布局函数\nfunction autoLayout(nodes, edges) {\n  // 用 dagre / elkjs 或自定义算法计算坐标\n  // 支持：思维导图放射布局、流程图垂直布局、网格布局\n}')
add_p('收益：')
add_p('- prompt 可以砍掉 ~40 行布局规则，更简洁\n- 坐标永远正确，不依赖 LLM 的"数学能力"\n- 支持更复杂的布局（如 10 分支思维导图）')

add_h3('改进二：Schema 驱动 JSON 校验 + 自动修复')
add_p('当前问题：parseAIResponse 只做了基本字段校验（type、actions、question），没有对每个 action 内部的字段做验证。')
add_p('方案：定义完整的 JSON Schema，解析后自动修复常见错误。')
add_code('// schema 定义\nconst ACTION_SCHEMA = {\n  add_shape: { required: [\'type\',\'shape\',\'x\',\'y\',\'width\',\'height\',\'id\'],\n               defaults: { backgroundColor: \'#4A90D9\', strokeColor: \'#000000\' },\n               range: { x: [50, 1150], y: [50, 750] } },\n  add_arrow: { required: [\'type\',\'fromId\',\'toId\'], autoFix: (a, actions) => {\n    // 如果 fromId/toId 不存在于 actions 中，尝试模糊匹配\n  }},\n  // ...\n};\n\nfunction validateAndFix(parsed) {\n  for (const action of parsed.actions) {\n    const schema = ACTION_SCHEMA[action.type];\n    // 补默认值、裁剪越界坐标、修复引用\n  }\n}')
add_p('收益：即使 AI 输出的 JSON 有小错误（坐标越界、缺少字段），前端能自动修正，不抛异常。')

add_h3('改进三：两段式生成（Plan → Execute）')
add_p('当前问题：一次 API 调用同时做"理解意图 + 生成坐标"，对复杂图表（思维导图、大流程图）来说 token 压力大，容易出错。')
add_p('方案：拆成两次轻量调用。')
add_code('第1次调用（Plan，轻量）：\n  Prompt 只包含：身份 + 核心原则 + 决策框架 + 简短示例\n  输出：{"type":"plan","nodes":[{label,shape,level,parent}],"edges":[...]}\n\n第2次调用（Execute，带坐标）：\n  Prompt：布局规则 + 颜色规则 + 自检清单\n  输入：plan + 前端算好的坐标建议\n  输出：{"type":"execute","actions":[...]}')
add_p('收益：\n- 每次调用的 prompt 更短，模型更容易遵循\n- Plan 阶段不需要关心坐标（前端算），Execute 阶段只做格式转换\n- 如果 Plan 正确但 Execute 失败，可以只重试 Execute')
add_p('代价：多一次 API 调用，延迟增加 1-2 秒。')

add_h3('改进四：Prompt 单元测试框架')
add_p('当前问题：改 prompt 后只能手动测试，不知道改一行会不会破坏其他 case。')
add_p('方案：建一个测试集，自动化回归。')
add_code('// prompt.test.js\nconst TEST_CASES = [\n  { input: \'画一个红色矩形\', expected: { type: \'execute\', actionCount: 1, firstAction: \'add_shape\' } },\n  { input: \'画西游记主要人物思维导图\', expected: { type: \'execute\', minActions: 9 } },\n  { input: \'你好\', expected: { type: \'clarify\' } },\n  { input: \'把红色改成蓝色\', canvasState: [...], expected: { type: \'execute\', firstAction: \'modify_shape\' } },\n  // ... 20+ cases\n];\n\nfor (const tc of TEST_CASES) {\n  const result = await callFullPipeline(tc.input, tc.canvasState);\n  assert(result.type === tc.expected.type);\n}')
add_p('收益：每次改 prompt 后跑一次测试，10 秒知道有没有回归。')

add_h3('改进五：分离 prompt 到独立文件')
add_p('当前问题：180 行的模板字符串嵌在 JS 文件里，不好 diff、不好版本管理。')
add_p('方案：把系统提示词抽到 .md 或 .txt 文件，构建时注入。')
add_code('// promptBuilder.js\nimport systemPrompt from \'./systemPrompt.txt?raw\';\nexport function buildSystemPrompt() { return systemPrompt; }')
add_p('收益：prompt 改动在 git diff 中一目了然，也可以用 markdown 预览查看结构。')

# ── 优先级 ──
add_h2('五、优先级建议')
add_table(
    ['优先级', '改进', '难度', '效果'],
    [
        ['P0', '前端自动布局', '中', '根本解决思维导图错位'],
        ['P1', 'Schema 校验 + 自动修复', '低', '减少 JSON 解析失败'],
        ['P1', 'Prompt 测试框架', '中', '防止改坏现有功能'],
        ['P2', '两段式生成', '中高', '提升复杂图表质量'],
        ['P2', 'prompt 文件分离', '低', '改善开发体验'],
    ]
)
add_p('最核心的问题是：LLM 不擅长生成精确坐标。只要把布局计算从 prompt 里剥离出来，让前端代码算坐标，系统提示词可以精简 40%，JSON 格式错误率也会同步下降。这是投入产出比最高的改进方向。')

# ═══════════════════ PART 2 ═══════════════════
doc.add_page_break()
add_h1('第二部分：Claude Code 提示词工程建议')
add_p('以下内容来自 05_ClaudeCode_PromptEngineering.md，完整原样保留。')

add_h2('任务：实现并验证指令解析层的完整Prompt工程体系')

add_h3('背景说明')
add_p('这是一款AI语音绘图工具。用户用语音说话，AI需要把自然语言转化为Excalidraw画布操作的JSON指令。')
add_p('核心挑战：这不是一个简单的"写一段system prompt"任务，而是一个需要精确设计的多层级Prompt体系。过去我们讨论认为这是整个项目技术含量最高的部分，原因是：')
add_p('1. 输出格式必须100%是合法JSON，任何多余文字都会导致解析失败\n2. 画布状态是动态的，每次调用内容不同，必须正确注入\n3. 中文自然语言的歧义性高，代词引用（"那个"）和模糊指令（"太乱了"）需要精确的规则\n4. clarify vs execute的判断边界不清晰时会造成体验崩溃')

add_h2('你需要实现的文件')

add_h3('文件1：frontend/src/utils/promptBuilder.js')
add_p('负责在每次调用前动态构建完整的消息结构。')

add_code('''/**
 * Prompt层级结构说明：
 *
 * [System Prompt] - 固定内容，每次调用完全相同
 *   ├── 角色定义
 *   ├── 输出格式约束（JSON Schema）
 *   ├── 颜色映射表
 *   ├── 坐标规则
 *   ├── 代词引用规则
 *   └── clarify触发规则
 *
 * [User Message] - 动态内容，每次调用前构建
 *   ├── ## 当前画布状态（来自canvasSerializer）
 *   ├── ## 对话历史（最近5轮，格式化为文字）
 *   └── ## 用户指令（本次语音转文字结果）
 *
 * 重要：画布状态和对话历史必须放在User Message，不放在System Prompt
 * 原因：System Prompt会被缓存复用，动态内容放进去会破坏缓存并浪费token
 */

export function buildSystemPrompt() {
  return `你是一个绘图指令解析器。你的唯一任务是将用户的自然语言指令转化为结构化的JSON操作指令。

## 严格输出约束
你必须只返回合法JSON对象，不包含任何其他文字、代码块标记（不要用\`\`\`）、解释或换行前缀。
第一个字符必须是 {，最后一个字符必须是 }。

## 输出格式

执行类型（当指令明确可执行时）：
{"type":"execute","actions":[...actions数组...],"confirmMessage":"一句话告知用户你做了什么"}

澄清类型（当指令模糊、目标不明确时）：
{"type":"clarify","question":"一句话问用户","options":["选项A","选项B","其他"]}

## Actions格式定义

add_shape（新增图形）：
{"type":"add_shape","shape":"rect|ellipse|text","x":数字,"y":数字,"width":数字,"height":数字,"backgroundColor":"hex色值","strokeColor":"#000000","label":"可选文字","id":"shape_时间戳"}

modify_shape（修改已有图形）：
{"type":"modify_shape","targetId":"元素id","changes":{"backgroundColor":"hex","x":数字,"y":数字,"width":数字,"height":数字,"label":"文字"}}

delete_shape（删除图形）：
{"type":"delete_shape","targetId":"元素id"}

add_arrow（连接两个图形）：
{"type":"add_arrow","fromId":"起点元素id","toId":"终点元素id","label":"可选文字","id":"arrow_时间戳"}

clear_canvas（清空画布）：
{"type":"clear_canvas"}

## 坐标系规则
画布尺寸约1200x800，左上角为原点(0,0)。
位置参考：左边x≈100-300，中间x≈400-600，右边x≈700-1000；上方y≈100-250，中间y≈300-500，下方y≈550-700。
默认尺寸：矩形width:120,height:60；圆形width:80,height:80；文字不需要width/height。
新增图形时，若用户未指定位置，自动选择画布空白区域，避免与已有元素重叠。

## 颜色关键词映射
红色→#FF4444, 蓝色→#4A90D9, 绿色→#52C41A, 黄色→#FADB14,
橙色→#FF7A00, 紫色→#7B5EA7, 黑色→#1A1A1A, 白色→#FFFFFF, 灰色→#8C8C8C,
粉色→#FF85C2, 青色→#13C2C2, 默认（无颜色指定）→#4A90D9

## 上下文引用规则
"那个"/"它"/"这个"→ 优先指对话历史中最近一次操作的元素id
"左边的"→ 画布状态中x值最小的同类元素
"右边的"→ 画布状态中x值最大的同类元素
"上面的"→ 画布状态中y值最小的同类元素
"刚才的"/"刚刚的" → 对话历史中最近一次add_shape或modify_shape的targetId
"所有"/"全部" → 对画布中所有同类元素批量生成多个action

## clarify触发规则（以下情况必须返回clarify，不得猜测执行）
1. 纯评价性语言且无操作方向：如"感觉不好看"、"太乱了"、"不够清晰"、"看起来怪怪的"
2. 目标元素有歧义：画布上有2个以上同类元素，且用户用了"那个"等模糊代词
3. 指令方向不唯一：如"调整一下大小"（调大还是调小？多少？）
4. clarify的options必须是具体可操作的选项，不是泛化描述

## 边界情况处理
- 画布为空时用户说"删除那个"→ 返回clarify，question:"画布上还没有图形，你想先画什么？"，options:["画一个矩形","画一个圆形","画一个流程图"]
- 一句话包含多个操作→ actions数组中按顺序包含多个action
- 用户说"撤销"/"退一步"→ 返回 {"type":"execute","actions":[{"type":"undo"}],"confirmMessage":"已撤销上一步操作"}
- 用户说"清空"/"全删了"→ 返回 {"type":"execute","actions":[{"type":"clear_canvas"}],"confirmMessage":"已清空画布"}`;
}

export function buildUserMessage(canvasState, conversationHistory, userInput) {
  // conversationHistory格式：[{role:'user',content:'...'},{role:'assistant',content:'...'}]
  // 只取最近5轮（10条消息）
  const recentHistory = conversationHistory.slice(-10);

  const historyText = recentHistory.length > 0
    ? recentHistory.map(m => `${m.role === 'user' ? '用户' : 'AI'}：${m.content}`).join('\\n')
    : '（无历史）';

  return `## 当前画布状态\\n${canvasState}\\n\\n## 对话历史\\n${historyText}\\n\\n## 用户指令\\n${userInput}`;
}

export function buildMessages(canvasState, conversationHistory, userInput) {
  return [
    { role: 'user', content: buildUserMessage(canvasState, conversationHistory, userInput) }
  ];
}''')

add_h3('文件2：frontend/src/utils/promptValidator.js')
add_p('这个文件的用途是开发阶段的自我验证工具，不在生产环境运行。')

add_code('''/**
 * Prompt质量验证器
 * 用途：在开发调试阶段，对系统prompt进行结构性检查
 * 使用方式：在浏览器控制台运行 validatePromptQuality()
 */

import { buildSystemPrompt, buildUserMessage } from './promptBuilder';

// 检查1：System Prompt结构完整性
export function checkSystemPromptStructure() {
  const prompt = buildSystemPrompt();
  const required = [
    { key: '严格输出约束', check: prompt.includes('严格输出约束') },
    { key: 'execute格式', check: prompt.includes('"type":"execute"') },
    { key: 'clarify格式', check: prompt.includes('"type":"clarify"') },
    { key: 'add_shape定义', check: prompt.includes('add_shape') },
    { key: 'modify_shape定义', check: prompt.includes('modify_shape') },
    { key: '颜色映射', check: prompt.includes('#FF4444') },
    { key: '坐标规则', check: prompt.includes('1200x800') },
    { key: 'clarify触发规则', check: prompt.includes('clarify触发规则') },
  ];

  const missing = required.filter(r => !r.check).map(r => r.key);
  if (missing.length === 0) {
    console.log('✅ System Prompt结构完整');
  } else {
    console.error('❌ System Prompt缺少以下部分：', missing);
  }
  return missing.length === 0;
}

// 检查2：User Message是否包含动态内容
export function checkUserMessageDynamic() {
  const canvas1 = '画布元素：[rect_001:矩形@(100,100),蓝色]';
  const canvas2 = '画布为空';
  const history = [{role:'user',content:'画一个矩形'},{role:'assistant',content:'已添加矩形'}];

  const msg1 = buildUserMessage(canvas1, history, '把那个变成红色');
  const msg2 = buildUserMessage(canvas2, [], '画一个圆');

  const checks = [
    { key: '画布状态被注入', check: msg1.includes('rect_001') },
    { key: '对话历史被注入', check: msg1.includes('画一个矩形') },
    { key: '用户指令被注入', check: msg1.includes('把那个变成红色') },
    { key: '空画布正确显示', check: msg2.includes('画布为空') },
    { key: '空历史正确显示', check: msg2.includes('无历史') },
  ];

  const failed = checks.filter(c => !c.check).map(c => c.key);
  if (failed.length === 0) {
    console.log('✅ User Message动态内容注入正确');
  } else {
    console.error('❌ User Message注入失败：', failed);
  }
  return failed.length === 0;
}

// 检查3：Token估算（粗略）
export function estimateTokenCount(canvasElements = 10, historyRounds = 5) {
  const mockCanvas = Array(canvasElements).fill('[rect_001:矩形@(100,100),蓝色,120x60]').join(' ');
  const mockHistory = Array(historyRounds * 2).fill({role:'user',content:'画一个矩形'});
  const mockInput = '把左边的圆变成红色并连接到右边的方块';

  const systemPrompt = buildSystemPrompt();
  const userMessage = buildUserMessage(mockCanvas, mockHistory, mockInput);

  // 粗略估算：中文1字≈1.5token，英文1词≈1token
  const systemTokens = Math.ceil(systemPrompt.length * 0.6);
  const userTokens = Math.ceil(userMessage.length * 0.6);
  const total = systemTokens + userTokens;

  console.log(`📊 Token估算（${canvasElements}个元素，${historyRounds}轮历史）：`);
  console.log(`   System Prompt: ~${systemTokens} tokens`);
  console.log(`   User Message: ~${userTokens} tokens`);
  console.log(`   总计: ~${total} tokens`);
  console.log(`   目标上限: 1000 tokens`);
  console.log(total <= 1000 ? '✅ 在目标范围内' : `⚠️ 超出目标 ${total - 1000} tokens，需要压缩`);

  return total;
}

export function validatePromptQuality() {
  console.log('=== Prompt质量检查 ===');
  checkSystemPromptStructure();
  checkUserMessageDynamic();
  estimateTokenCount(10, 5);
  estimateTokenCount(30, 5);  // 压力测试：30个元素时的token消耗
  console.log('=== 检查完成 ===');
}''')

add_h3('文件3：frontend/src/utils/promptTestCases.js')
add_p('开发阶段用于手动测试prompt效果的用例集，不在生产环境运行。')

add_code('''/**
 * 系统Prompt测试用例
 * 使用方式：开发时在控制台调用 runTestCase(案例ID) 查看AI返回
 * 验证AI返回是否符合预期
 */

export const TEST_CASES = [
  // --- P0：基础执行类 ---
  {
    id: 'TC001',
    desc: '基础：画矩形',
    canvasState: '画布为空',
    history: [],
    input: '画一个矩形',
    expect: { type: 'execute', mustContain: ['add_shape', 'rect'] }
  },
  {
    id: 'TC002',
    desc: '基础：画红色大圆在左上角',
    canvasState: '画布为空',
    history: [],
    input: '在左上角画一个红色大圆',
    expect: { type: 'execute', mustContain: ['ellipse', '#FF4444'], xRange: [50, 300], yRange: [50, 250] }
  },
  {
    id: 'TC003',
    desc: '基础：连接两个图形',
    canvasState: '画布元素：[rect_001:矩形@(100,200),蓝色,120x60] [rect_002:矩形@(500,200),绿色,120x60]',
    history: [],
    input: '用箭头连接左边的方块和右边的方块',
    expect: { type: 'execute', mustContain: ['add_arrow', 'rect_001', 'rect_002'] }
  },

  // --- P0：上下文引用类 ---
  {
    id: 'TC004',
    desc: '上下文：代词引用"那个"',
    canvasState: '画布元素：[rect_001:矩形@(100,100),蓝色,120x60]',
    history: [
      { role: 'user', content: '画一个蓝色矩形' },
      { role: 'assistant', content: '{"type":"execute","actions":[{"type":"add_shape","id":"rect_001"...}],"confirmMessage":"已添加蓝色矩形"}' }
    ],
    input: '把那个变成红色',
    expect: { type: 'execute', mustContain: ['modify_shape', 'rect_001', '#FF4444'] }
  },
  {
    id: 'TC005',
    desc: '上下文：位置代词"左边的"',
    canvasState: '画布元素：[rect_001:矩形@(100,200),蓝色] [rect_002:矩形@(500,200),绿色]',
    history: [],
    input: '删除左边的矩形',
    expect: { type: 'execute', mustContain: ['delete_shape', 'rect_001'] }
  },

  // --- P1：澄清类 ---
  {
    id: 'TC006',
    desc: '澄清：纯评价性语言',
    canvasState: '画布元素：[rect_001:矩形@(100,100),蓝色] [rect_002:矩形@(110,110),绿色]',
    history: [],
    input: '感觉不好看',
    expect: { type: 'clarify', mustHaveOptions: true, minOptions: 2 }
  },
  {
    id: 'TC007',
    desc: '澄清：代词歧义（多个同类元素）',
    canvasState: '画布元素：[rect_001:矩形@(100,100),蓝色] [rect_002:矩形@(400,100),绿色] [rect_003:矩形@(700,100),红色]',
    history: [],
    input: '把那个矩形变大',
    expect: { type: 'clarify', mustHaveOptions: true }
  },

  // --- 边界类 ---
  {
    id: 'TC008',
    desc: '边界：画布为空时说删除',
    canvasState: '画布为空',
    history: [],
    input: '删除那个圆',
    expect: { type: 'clarify' }  // 不应该返回execute
  },
  {
    id: 'TC009',
    desc: '边界：一句话多个操作',
    canvasState: '画布为空',
    history: [],
    input: '画一个红色矩形，再画一个蓝色圆，用箭头连起来',
    expect: { type: 'execute', minActionsCount: 3 }
  },
  {
    id: 'TC010',
    desc: '边界：带标签的图形',
    canvasState: '画布为空',
    history: [],
    input: '画一个矩形，里面写"用户登录"',
    expect: { type: 'execute', mustContain: ['用户登录'] }
  }
];''')

add_h2('实现要求')

add_h3('1. promptBuilder.js')
add_p('完整实现上方的三个导出函数，不得改变函数签名：')
add_p('- buildSystemPrompt() → 返回string\n- buildUserMessage(canvasState, conversationHistory, userInput) → 返回string\n- buildMessages(canvasState, conversationHistory, userInput) → 返回messages数组')

add_h3('2. promptValidator.js')
add_p('完整实现上方的验证函数，用于开发调试，不导出到生产bundle。')

add_h3('3. promptTestCases.js')
add_p('完整实现测试用例集，并额外实现一个运行函数：')

add_code('''export async function runTestCase(testCaseId, claudeApiCaller) {
  const tc = TEST_CASES.find(t => t.id === testCaseId);
  if (!tc) { console.error('找不到用例', testCaseId); return; }

  console.log(`\\n=== 运行测试用例 ${tc.id}：${tc.desc} ===`);
  console.log('输入：', tc.input);
  console.log('画布状态：', tc.canvasState);

  const result = await claudeApiCaller(tc.canvasState, tc.history, tc.input);

  console.log('AI返回：', JSON.stringify(result, null, 2));

  // 自动验证
  const pass = tc.expect.type === result.type;
  console.log(pass ? `✅ 类型正确（${result.type}）` : `❌ 类型错误，期望${tc.expect.type}，实际${result.type}`);

  if (tc.expect.mustContain) {
    const resultStr = JSON.stringify(result);
    tc.expect.mustContain.forEach(keyword => {
      const ok = resultStr.includes(keyword);
      console.log(ok ? `✅ 包含关键字：${keyword}` : `❌ 缺少关键字：${keyword}`);
    });
  }

  return result;
}''')

add_h3('4. 集成到 useCommandProcessor.js')
add_p('在 useCommandProcessor.js 中使用 promptBuilder.js：')

add_code('''import { buildSystemPrompt, buildMessages } from '../utils/promptBuilder';
import { serializeCanvas } from '../utils/canvasSerializer';

// 调用Claude API时：
const systemPrompt = buildSystemPrompt();
const canvasState = serializeCanvas(currentElements);
const messages = buildMessages(canvasState, conversationHistory, userInput);

// POST /api/chat 时传入：
{
  system: systemPrompt,
  messages: messages,
  userApiKey: sessionStorage.getItem('userApiKey') || ''
}''')

add_h2('验收标准')
add_p('开发完成后，请在浏览器控制台运行以下验证：')

add_code('''// 1. 结构检查
import { validatePromptQuality } from './utils/promptValidator';
validatePromptQuality();
// 期望：所有检查✅，token估算<1000

// 2. 手动跑测试用例（需要真实API连通后）
import { runTestCase } from './utils/promptTestCases';
// 至少跑 TC001、TC004、TC006、TC008 这4个核心用例
// TC001、TC004 必须返回 execute
// TC006、TC008 必须返回 clarify''')

add_h3('Day 1 Prompt层验收标准：')
add_p('- validatePromptQuality() 全部✅\n- TC001（基础画图）通过\n- TC004（上下文引用）通过\n- TC006（模糊指令澄清）通过')

add_h3('Day 2 Prompt层验收标准：')
add_p('- TC001-TC010 全部通过\n- 30个画布元素时token估算仍 < 1000')

# ═══════════════════ 保存 ═══════════════════
output_path = '/Users/zhangyuxiang/Downloads/VoiceDraw_系统提示词架构说明.docx'
doc.save(output_path)
print(f'已保存: {output_path}')
